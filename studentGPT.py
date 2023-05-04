import os
import re
import json
import time
import docx
import PyPDF2
import openai
import configparser
from colorama import Fore, Back, Style


def print_head():
    print(f"{Fore.CYAN}{Style.BRIGHT}========================================== {Fore.BLUE}Welcome to StudentGPT! {Fore.MAGENTA}{studentGPT_version}v{Fore.CYAN} ============================================={Style.RESET_ALL}")

def print_break_line():
    print(f"{Fore.CYAN}{Style.BRIGHT}==================================================================================================={Style.RESET_ALL}")

def estimate_reading_time(text):
    return len(text.split()) // 235

def list_of_task():
    TASKS = [
    "Summarize the paper",
    "Extract main points from the paper in numeric order",
    "Make revision notes from class notes",
    "Make a quiz out of the text",
    "Extract and define main keywords",
    "StudentGPT Settings",
    "StudentGPT Info",
    "Exit"
    ]
    print('\n'.join(f"{i+1}. {Back.YELLOW}{Fore.BLACK}{Style.BRIGHT}{task}{Style.RESET_ALL}" for i, task in enumerate(TASKS)))
    while True:
        user_input = input("Enter the task you want to proceed with: ")
        if user_input.isdigit() and 1 <= int(user_input) <= len(TASKS) and user_input != "8":
            return int(user_input)
        elif user_input == "8":
            print(f"{Back.CYAN}{Fore.BLACK}Bye Bye Now{Style.RESET_ALL}!")
            quit()
        else:
            print(f"{Back.RED}{Fore.BLACK}Invalid input. Please choose [1-{len(TASKS)}].{Style.RESET_ALL}")

def list_of_settings():
    while True:
        user_input = input("Enter the setting you want to change: ")
        if user_input == "1":
            global settings_output
            if settings_output == "Terminal":
                settings_output = "File"
                modify_config('main', 'output', "File")
            elif settings_output == "File":
                settings_output = "Both"
                modify_config('main', 'output', "Both")
            elif settings_output == "Both":
                settings_output = "Terminal"
                modify_config('main', 'output', "Terminal")
            return ""
        elif user_input == "0":
            return "exit"
        else:
            print(f"{Back.RED}{Fore.BLACK}Invalid input. Please choose [1-1].{Style.RESET_ALL}")

def get_user_choice(files):
    while True:
        try:
            choice = int(input("Enter the number of the file you want to proceed with: "))
            if choice == 0:
                return ""
            elif 1 <= choice <= len(files):
                return files[choice - 1]
            else:
                print("Invalid input. Please enter a valid number from the list.")
        except ValueError:
            print("Invalid input. Please enter a number.")

def split_text(text, max_length):
    chunks = []
    start = 0
    end = max_length
    while start < len(text):
        if end < len(text):
            end = text.rfind('.', start, end) + 2
        else:
            end = len(text)
        chunks.append(text[start:end])
        start = end
        end = start + max_length
    return chunks

def extract_text(pdf_file):
    with open(pdf_file, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        text = ''
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            try:
                page_text = page.extract_text()
            except Exception as e:
                print(f"{Back.RED}{Fore.BLACK}An error occurred while extracting text from page " + Back.CYAN + f" {page_num} {Back.RED}: {e}" + Style.RESET_ALL)
                print(f"{Back.RED}{Fore.BLACK}The page was skipped and it won't be used for anything! Try and generate new PDF by copying all the text into Microsoft Word > Export to PDF and then use the new PDF here.{Style.RESET_ALL}")
                continue
            text += page_text
    return text

def clean_text(text):
    # Remove page numbers (assuming they are at the beginning or end of a line)
    text = re.sub(r'^\d+|\d+$', '', text, flags=re.MULTILINE)
    # Remove headers and footers (assuming they are in uppercase)
    text = re.sub(r'^[A-Z\s]+$', '', text, flags=re.MULTILINE)
    # Remove extra whitespaces
    text = re.sub(r'\s+', ' ', text).strip()
    # Remove links
    text = re.sub(r"http\S+", "", text)
    return text

def read_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)

def generate_revision_notes(text):
    completion = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "user", "content": "Create extensive revision notes from the following notes: " + text}
            ]
            )
    completion_str = json.dumps(completion.choices[0].message, indent=2)
    response = json.loads(completion_str)
    return response["content"]

def save_to_file(text_chunks, file_name, task):
    chosen_output_file = ".docx"
    if task == "summary":
        final_file_name = re.sub(r'[^a-zA-Z0-9 ]', '', file_name)
        chosen_output_file = f"Summary of {final_file_name}" + chosen_output_file
        file_path = os.path.join(os.path.dirname(__file__), chosen_output_file)
        if os.path.isfile(file_path):
            print(f"{Back.RED}{Fore.BLACK} The file '{Back.WHITE}{file_name}{Back.RED}' already exists in the current directory! Couldn't save it :( {Style.RESET_ALL}")
            return False
        else:
            output_file = docx.Document()
            output_file.add_heading(f'Summarization of {file_name}', level=1)
            for i, paragraph in enumerate(text_chunks):
                    output_file.add_paragraph(text_chunks[i] + "\n")
            output_file.save(chosen_output_file)
    elif task == "extraction_main_points":
        final_file_name = re.sub(r'[^a-zA-Z0-9 ]', '', file_name)
        chosen_output_file = f"Main points of {final_file_name}" + chosen_output_file
        file_path = os.path.join(os.path.dirname(__file__), chosen_output_file)
        if os.path.isfile(file_path):
            print(f"{Back.RED}{Fore.BLACK} The file '{Back.WHITE}{file_name}{Back.RED}' already exists in the current directory! Couldn't save it :( {Style.RESET_ALL}")
            return False
        else:
            output_file = docx.Document()
            output_file.add_heading(f'Main points of {file_name}', level=1)
            for i, paragraph in enumerate(text_chunks):
                    output_file.add_paragraph(text_chunks[i] + "\n")
            output_file.save(chosen_output_file)
    elif task == "quiz":
        final_file_name = re.sub(r'[^a-zA-Z0-9 ]', '', file_name)
        chosen_output_file = f"Quiz made from {final_file_name}" + chosen_output_file
        file_path = os.path.join(os.path.dirname(__file__), chosen_output_file)
        if os.path.isfile(file_path):
            print(f"{Back.RED}{Fore.BLACK} The file '{Back.WHITE}{file_name}{Back.RED}' already exists in the current directory! Couldn't save it :( {Style.RESET_ALL}")
            return False
        else:
            output_file = docx.Document()
            output_file.add_heading(f'Quiz made from {file_name}', level=1)
            for i, paragraph in enumerate(text_chunks):
                    output_file.add_paragraph(text_chunks[i] + "\n")
            output_file.save(chosen_output_file)
    elif task == "key_words":
        final_file_name = re.sub(r'[^a-zA-Z0-9 ]', '', file_name)
        chosen_output_file = f"Key definitions of {final_file_name}" + chosen_output_file
        file_path = os.path.join(os.path.dirname(__file__), chosen_output_file)
        if os.path.isfile(file_path):
            print(f"{Back.RED}{Fore.BLACK} The file '{Back.WHITE}{file_name}{Back.RED}' already exists in the current directory! Couldn't save it :( {Style.RESET_ALL}")
            return False
        else:
            output_file = docx.Document()
            output_file.add_heading(f'Key definitions of {file_name}', level=1)
            for i, paragraph in enumerate(text_chunks):
                    output_file.add_paragraph(text_chunks[i] + "\n")
            output_file.save(chosen_output_file)


    print(f"The output was succesfully saved to {Back.WHITE}{Fore.BLACK}{chosen_output_file}{Style.RESET_ALL}")
    return True

def print_text_file_info(filename, filepath, clean_text):
    print_break_line()
    print("Chosen: " + Fore.BLACK + Back.BLUE + Style.BRIGHT + filename + Style.RESET_ALL)
    print("Text extracted from " + Back.WHITE + Fore.BLACK + f"{filepath}" + Style.RESET_ALL )
    reading_time = estimate_reading_time(clean_text)
    print(f"Estimated reading time: {Back.WHITE}{Fore.BLACK}{reading_time}{Style.RESET_ALL} minutes.")
    print(f"Total count: " + Back.WHITE + Fore.BLACK + str(len(clean_text)) + Style.RESET_ALL + " characters")
    print_break_line()

def clean_terminal():
    os.system('cls')

def modify_config(section, key, value):
    config[section] = {key: value}
    with open(config_file, 'w') as configfile:
        config.write(configfile)







### MAIN START
studentGPT_version = 0.30
settings_output = "Terminal"
config_file = 'config.ini'
clean_terminal()

# Loading config
config = configparser.ConfigParser()
if os.path.exists(config_file):
    config.read(config_file)
    try:
        settings_output = config.get('main', 'output')
        settings_api_key = config.get('auth', 'openai_api_key')
        openai.api_key = settings_api_key
    except (configparser.NoSectionError, configparser.NoOptionError) as e:
        print(f"{Back.RED}{Fore.BLACK}An error occurred while reading config file: {e}. \nRemove the 'config.ini' file and restart the program.{Style.RESET_ALL}")
        quit()
else:
    settings_api_key = ""
    print("First time start detected. Please enter your OpenAI API key now. (You can get it from OpenAI account settings).")
    while len(settings_api_key) != 51:
        settings_api_key = input("Please enter a 51-character-long OpenAI API key: ")
        if len(settings_api_key) != 51:
            print("The input must be exactly 51 characters long. Please try again.")
    modify_config('main', 'output', settings_output)
    modify_config('auth', 'openai_api_key', settings_api_key)
    openai.api_key = settings_api_key
    clean_terminal()



# Validating the config file
if(settings_output != "Terminal" and settings_output != "Both" and settings_output != "File"):
        print(f"{Back.RED}{Fore.BLACK}An error occurred while reading config file: 'output' value is invalid. \nFixing it...{Style.RESET_ALL}")
        modify_config('main', 'output', "Terminal")

#Checking API key
try:
    response = openai.Completion.create(prompt="Hello world", model="text-davinci-003")
except openai.error.APIError as e:
    print(f"{Back.RED}{Fore.BLACK}OpenAI API returned an API Error: {e}{Style.RESET_ALL}")
    quit()
except openai.error.APIConnectionError as e:
    print(f"{Back.RED}{Fore.BLACK}Failed to connect to OpenAI API: {e}{Style.RESET_ALL}")
    quit()
except openai.error.RateLimitError as e:
    print(f"{Back.RED}{Fore.BLACK}OpenAI API request exceeded rate limit: {e}{Style.RESET_ALL}")
    quit()
except openai.error.AuthenticationError as e:
    print(f"{Back.RED}{Fore.BLACK}OpenAI API key is invalid: {e}{Style.RESET_ALL}")
    quit()








while True:
    print_head()
    chosen_task = list_of_task()
    clean_terminal()
    if chosen_task == 1 or chosen_task == 2 or chosen_task == 4 or chosen_task == 5:
        print_head()
        pdf_files = [f for f in os.listdir('Readings') if f.endswith('.pdf')]
        print("List of PDF files:")
        print(f"{0}. {Fore.BLACK}{Back.MAGENTA}{Style.BRIGHT}Go back to previous menu.{Style.RESET_ALL}")
        for i, pdf_file in enumerate(pdf_files, start=1):
            print(f"{i}. {Fore.BLACK}{Back.YELLOW}{Style.BRIGHT}{pdf_file}{Style.RESET_ALL}")
        chosen_pdf_file = get_user_choice(pdf_files)
        if(chosen_pdf_file == ""):
            clean_terminal()
            continue

        # Showing the chosen file info
        pdf_file_path = os.path.join('Readings', chosen_pdf_file)
        extracted_text = extract_text(pdf_file_path)
        cleaned_text = clean_text(extracted_text)
        print_text_file_info(chosen_pdf_file, pdf_file_path, cleaned_text)


        if chosen_task == 1:
            print(f"Beginning summarisation")
            print_break_line()
            split_articles = split_text(cleaned_text, 14000)
            final_result = []
            for idx, article_indv in enumerate(split_articles):
                print(f"Part {Back.WHITE}{Fore.BLACK}{idx + 1}/" + str(len(split_articles)) + f"{Style.RESET_ALL}:\n")
                completion = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "user", "content": "Summarise the article in a longer paragraph: " + article_indv}
                ]
                )
                completion_str = json.dumps(completion.choices[0].message, indent=2)
                response = json.loads(completion_str)
                final_result.append(response["content"])
                if settings_output == "Terminal" or settings_output == "Both":
                    print(Fore.WHITE + Style.BRIGHT + response["content"] + Style.RESET_ALL)
                print_break_line()
                time.sleep(8)
            if settings_output == "File" or settings_output == "Both":
                save_to_file(final_result, chosen_pdf_file, "summary")

        elif chosen_task == 2:

            print(f"Beginning extraction of the main points")
            print_break_line()
            split_articles = split_text(cleaned_text, 14000)
            final_result = []
            for idx, article_indv in enumerate(split_articles):
                print(f"Part {Back.WHITE}{Fore.BLACK}{idx + 1}/" + str(len(split_articles)) + f"{Style.RESET_ALL}:\n")
                completion = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "user", "content": "Highlight important facts in this article in numeric order. Make sure to actually explain the facts rather than saying that the article mentions them. The Article: " + article_indv}
                ]
                )
                completion_str = json.dumps(completion.choices[0].message, indent=2)
                response = json.loads(completion_str)
                non_empty_lines = [line for line in response["content"].splitlines() if line.strip()]
                resulting_points = "\n".join(non_empty_lines)
                final_result.append(resulting_points)
                if settings_output == "Terminal" or settings_output == "Both":
                    print(Fore.WHITE + Style.BRIGHT + resulting_points + Style.RESET_ALL)
                print_break_line()
                time.sleep(8)
            if settings_output == "File" or settings_output == "Both":
                save_to_file(final_result, chosen_pdf_file, "extraction_main_points")

                

        elif chosen_task == 4:
            split_articles = split_text(cleaned_text, 5000)
            print(f"{Back.RED}{Fore.BLACK}Warning! {Style.RESET_ALL} - You are using AI to generate the quiz. AI is not perfect and there is a possibility of AI asking wrong questions / giving wrong answers.")
            print_break_line()
            answers = {}
            final_result = []
            for idx, article_indv in enumerate(split_articles):
                print(f"Part {Back.WHITE}{Fore.BLACK}{idx + 1}/" + str(len(split_articles)) + f"{Style.RESET_ALL}:\n")
                completion = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                #model="gpt-4",
                messages=[
                    {"role": "user", "content": "Please make me a multiple choice quiz based on the given revision notes. Make sure each question has 4 answers (one correct and three made up). Make sure the questions are in numeric order. Make sure the order of the four answers is randomised so the correct answer is not always the same. Give me the answer for each question after you list all the possible answers for that specific question in a format: 'Answer: ANSWER_LETTER_HERE' etc. The Notes: " + article_indv}
                ]
                )
                completion_str = json.dumps(completion.choices[0].message, indent=2)
                response = json.loads(completion_str)
                non_empty_lines = [line for line in response["content"].splitlines() if line.strip()]
                resulting_points = "\n".join(non_empty_lines)
                index = 0
                for line in resulting_points.splitlines():
                    if line.strip() and line.strip()[0].isdigit():
                        if settings_output == "Terminal" or settings_output == "Both":
                            print(f"{Back.BLUE}{Fore.BLACK}{line}{Style.RESET_ALL}")
                        final_result.append(line)
                    elif "Answer" in line:
                        if settings_output == "Terminal" or settings_output == "Both":
                            print(f"{Back.GREEN}{Fore.BLACK}{line}{Style.RESET_ALL}")
                        final_result.append(line)
                    else:
                        if settings_output == "Terminal" or settings_output == "Both":
                            print(line)
                        final_result.append(line)
                    index += 1
                print_break_line()
                time.sleep(8)
            if settings_output == "File" or settings_output == "Both":
                save_to_file(final_result, chosen_pdf_file, "quiz")


        elif chosen_task == 5:
            split_articles = split_text(cleaned_text, 14000)
            keywords = []
            final_result = []
            for idx, article_indv in enumerate(split_articles):
                print("Part " + Back.WHITE + Fore.BLACK + f"{idx + 1}/" + str(len(split_articles)) + Style.RESET_ALL + ":\n")
                if(idx == 0):
                    completion = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "user", "content": "Extract key terms or words from the text which the text is talking about and put them in this format, separated by comas: 'KEY_WORD_HERE, KEY_WORD_HERE'. Make sure to only extract important ones. Don't give me paragraphs of text, or their explanations only in the format I requested. Make sure not to repeat the same key words. The Text: " + article_indv}
                    ]
                    )
                else:
                    keywords_strings = ", ".join(keywords)
                    completion = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "user", "content": "Extract key terms or words from the text which the text is talking about and put them in this format, separated by comas: 'KEY_WORD_HERE, KEY_WORD_HERE'. Make sure to only extract important ones. Don't give me paragraphs of text, or their explanations only in the format I requested. Make sure not to repeat the same key words. If these words or key terms come up, ignore them: " + keywords_strings + ". The Text: " + article_indv}
                    ]
                    )
                completion_str = json.dumps(completion.choices[0].message, indent=2)
                response = json.loads(completion_str)
                keywords_extracted = response["content"].split(", ")
                for keyword in keywords_extracted:
                    if keyword.lower() not in [x.lower() for x in keywords]:
                        completion = openai.ChatCompletion.create(
                        model="gpt-3.5-turbo",
                        messages=[
                            {"role": "user", "content": "Give me a short (max two sentences) and easy to understand definition of " + keyword}
                        ]
                        )
                        completion_str = json.dumps(completion.choices[0].message, indent=2)
                        response = json.loads(completion_str)
                        if settings_output == "Terminal" or settings_output == "Both":
                            print(Back.CYAN + Fore.BLACK + keyword.capitalize() + Style.RESET_ALL + " --> " + response["content"]) 
                            print("-----------")
                        final_result.append(keyword.capitalize() + " --> " + response["content"])
                        keywords.append(keyword)
                print_break_line()
                time.sleep(8)
            if settings_output == "File" or settings_output == "Both":
                save_to_file(final_result, chosen_pdf_file, "key_words")

    elif chosen_task == 3:
        print_head()
        notes_folder = 'Notes'
        output_file = 'revision_notes.docx'
        combined_notes_doc = docx.Document()
        idx = 0
        docx_files = [f for f in os.listdir('Notes') if f.endswith('.docx')]
        docx_files.insert(0, "All of them")
        print("List of Notes:")
        print(f"{0}. {Back.MAGENTA}{Fore.BLACK}{Style.BRIGHT}Go back to previous menu.{Style.RESET_ALL}")
        for i, docx_file in enumerate(docx_files, start=1):
            print(f"{i}. " + Back.YELLOW  + Fore.BLACK + Style.BRIGHT + f"{docx_file}" + Style.RESET_ALL)
        chosen_docx_file = get_user_choice(docx_files)
        if(chosen_docx_file == ""):
            clean_terminal()
            continue
        elif(chosen_docx_file == "All of them"):
            for file in os.listdir(notes_folder):
                print("Making notes out of " + Back.WHITE + Fore.BLACK + f"{file}" + Style.RESET_ALL + ". Part " + Back.WHITE + Fore.BLACK + f"{idx + 1}/" + str(len(os.listdir(notes_folder))) + Style.RESET_ALL + ".")
                if file.endswith('.docx'):
                    file_path = os.path.join(notes_folder, file)
                    text = read_docx(file_path)
                    revision_notes = generate_revision_notes(text)
                    combined_notes_doc.add_heading(file, level=1)
                    combined_notes_doc.add_paragraph(revision_notes)
                    combined_notes_doc.add_page_break()
                    idx = 1 + idx
                    time.sleep(8)
                    print(Back.WHITE + Fore.BLACK + "Done!" + Style.RESET_ALL + "\n")
        else:
            print("Making notes out of " + Back.WHITE + Fore.BLACK + f"{chosen_docx_file}" + Style.RESET_ALL )
            file_path = os.path.join(notes_folder, chosen_docx_file)
            text = read_docx(file_path)
            revision_notes = generate_revision_notes(text)
            combined_notes_doc.add_heading(chosen_docx_file, level=1)
            combined_notes_doc.add_paragraph(revision_notes)
            combined_notes_doc.add_page_break()
            print(Back.WHITE + Fore.BLACK + "Done!" + Style.RESET_ALL + "\n")
        combined_notes_doc.save(output_file)
        print(f"Revision notes saved to " + Back.WHITE + Fore.BLACK + f"{output_file}" + Style.RESET_ALL)
    elif chosen_task == 6:
        while True:
            clean_terminal()
            print_head()
            print(f"0. {Fore.BLACK}{Back.MAGENTA}{Style.BRIGHT}Go back to previous menu{Style.RESET_ALL}")
            output_options = {
                "Terminal": f"1. Output of the tasks: {Back.GREEN}{Fore.BLACK}Terminal{Style.RESET_ALL} / File / Both",
                "File": f"1. Output of the tasks: Terminal / {Back.GREEN}{Fore.BLACK}File{Style.RESET_ALL} / Both",
                "Both": f"1. Output of the tasks: Terminal / File / {Back.GREEN}{Fore.BLACK}Both{Style.RESET_ALL}"
            }
            print(output_options.get(settings_output, f"Output of the tasks: {Back.MAGENTA}{Fore.BLACK}Terminal{Style.RESET_ALL} / File / Both"))
            results = list_of_settings()
            if results == "exit":
                break
            elif results == "":
                continue
    elif chosen_task == 7:
        clean_terminal()
        print_head()
        info_string = '''StudentGPT is a powerful program designed to assist students in their studies using the ChatGPT API. It was developed as a side project to help me summarize longer readings and I thought it would be cool to expand it further with more features. The program provides a range of tools and features to help students understand and remember the information they are studying.

StudentGPT can summarize papers to provide an overview of key points, generate quizzes to test students' knowledge and identify weak areas, extract key words and definitions to expand vocabulary as well as generate revision notes based on the in-class-notes.

Overall, StudentGPT is an invaluable tool for any student looking to improve their understanding and retention of the information they are studying. Whether you're preparing for an exam or simply trying to keep up with your coursework, StudentGPT can help you achieve your academic goals.'''
        print(f"{Fore.GREEN}{info_string}{Style.RESET_ALL}")
        print("\n")
        print(f"Made by {Back.MAGENTA}{Fore.BLACK} michpcx {Style.RESET_ALL}")
        print(f"Current version: {Back.MAGENTA}{Fore.BLACK} {studentGPT_version}v {Style.RESET_ALL}")
        print(f"Github link: {Back.MAGENTA}{Fore.BLACK} https://github.com/michpcx/StudentGPT {Style.RESET_ALL}")
        print("\n")
    print(Back.GREEN + Fore.BLACK + "I hope that helps! :) " + Style.RESET_ALL + "\n \n")